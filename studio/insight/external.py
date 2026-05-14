"""External reports + GPT-4o integration.

Users sometimes have analysis reports from elsewhere — consulting decks,
competitor teardowns, things they had ChatGPT generate, etc. Letting them
upload those and have the tool reference them downstream is far more valuable
than re-running our own analysis when they already trust an outside source.

Flow:
    1. Upload text (paste / .md / .txt) → studio_external_reports row.
    2. Integrate button → gpt-4o reads all uploaded reports (and optionally
       the latest tool-generated consensus) → emits a consensus-shaped JSON
       that the existing InsightReport renderer can display unchanged.
    3. Strategy / Composer prompts now include both:
        - tool-generated consensus (if any), and
        - any integrated report (if any) — both contribute to downstream.
"""
from __future__ import annotations

import json
import time
import uuid
from typing import Any

from .. import db, project
from ..generators import registry
from ..llm_call import call_for_json


# ---- Upload / list / delete ---------------------------------------------

def extract_text_from_bytes(filename: str, data: bytes) -> tuple[str, str, str | None]:
    """Best-effort text extraction from any uploaded file.

    Returns: (extracted_text, detected_format, error_or_none)

    Strategy by extension:
      .pdf            → pypdf
      .docx           → python-docx
      .tex/.md/.txt/  → utf-8 decode (with errors='replace' fallback)
      otherwise       → try utf-8 first, then latin-1, save whatever decodes
    """
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    # PDF
    if ext == "pdf":
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(data))
            pages: list[str] = []
            for i, page in enumerate(reader.pages):
                try:
                    pages.append(page.extract_text() or "")
                except Exception as pe:
                    pages.append(f"[第 {i+1} 页解析失败：{pe}]")
            text = "\n\n".join(pages).strip()
            if not text:
                return ("", "pdf",
                        "PDF 里读不到文字（可能是扫描件 / 图片型 PDF）。"
                        "建议导出成文字版后再上传，或者直接粘贴文本。")
            return (text, "pdf", None)
        except ImportError:
            return ("", "pdf",
                    "服务器还没装 pypdf — 后端 requirements 没更新到位。")
        except Exception as e:
            return ("", "pdf", f"PDF 解析失败：{e}")

    # DOCX
    if ext == "docx":
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(data))
            # v0.59.5: python-docx 的 doc.paragraphs 只读 body-level 段落，
            # 完全跳过表格内容。产品说明 / 人设卡 / 功能清单这类文档大量
            # 用表格组织信息 → 之前提取出来空字符串，上传报 400。
            # 现在递归扫表格（含嵌套表格 + 表格里的段落）。
            chunks: list[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    chunks.append(p.text.strip())

            def _walk_tables(tables) -> None:
                for tbl in tables:
                    for row in tbl.rows:
                        cells_text: list[str] = []
                        for cell in row.cells:
                            cell_parts: list[str] = []
                            for p in cell.paragraphs:
                                if p.text.strip():
                                    cell_parts.append(p.text.strip())
                            # Nested tables inside this cell — rare but happens
                            if cell.tables:
                                _walk_tables(cell.tables)
                            if cell_parts:
                                cells_text.append(" ".join(cell_parts))
                        if cells_text:
                            chunks.append(" | ".join(cells_text))

            _walk_tables(doc.tables)
            text = "\n\n".join(chunks).strip()
            if not text:
                return ("", "docx",
                        "Word 文档里没读到任何文本（段落 + 表格都为空 — "
                        "可能是纯图片 / 纯嵌入对象，建议导出为文字版后再上传）。")
            return (text, "docx", None)
        except ImportError:
            return ("", "docx", "服务器没装 python-docx。")
        except Exception as e:
            return ("", "docx", f"DOCX 解析失败：{e}")

    # Everything else — try plain text decode. Includes .tex, .md, .markdown,
    # .txt, .html, etc. The user said: 不设格式都行，给 AI 读。
    fmt = ext or "text"
    for enc in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            text = data.decode(enc)
            return (text, fmt, None)
        except UnicodeDecodeError:
            continue
    # Last resort
    return (data.decode("utf-8", errors="replace"), fmt,
            "文件不是常见的文本编码，已尽量解码（可能有乱码）。")


def save_external_report(
    *, name: str, content: str,
    library_id: str | None = None,
    source: str = "粘贴文本",
    format: str = "text",
) -> dict[str, Any]:
    db.apply_migrations(verbose=False)
    project.ensure_bootstrap()
    pid = project.active_project_id()
    report_id = "ext_" + uuid.uuid4().hex[:14]
    now = int(time.time())
    with db.connect() as con:
        con.execute(
            "INSERT INTO studio_external_reports"
            " (report_id, project_id, library_id, name, source, format, content, uploaded_at)"
            " VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (report_id, pid, library_id, name, source, format, content, now),
        )
    # Drop the downstream reference-block cache so the new report is
    # visible immediately to Strategy/Composer prompts.
    try:
        from . import pipeline as _ip
        _ip.invalidate_ref_block_cache()
    except Exception:
        pass
    return {
        "report_id": report_id, "project_id": pid, "library_id": library_id,
        "name": name, "source": source, "format": format,
        "content_chars": len(content), "uploaded_at": now,
    }


def list_external_reports(library_id: str | None = None) -> list[dict[str, Any]]:
    """v0.61.5 ：studio_external_reports 也是 per-library .db scoped。
    用户上传 .db 后 active_lib 会切到新库，老库里的 external_reports 行从
    current_db_path 看就消失了。修复 ：扫所有 data/libraries/*/xhs.db，
    union 当前项目下的 external_reports 行 + 按 uploaded_at DESC 排序。"""
    import sqlite3
    from .. import config, library as _lib

    db.apply_migrations(verbose=False)
    project.ensure_bootstrap()
    pid = project.active_project_id()

    libs_dir = config.DATA_DIR / "libraries"
    seen: dict[str, dict[str, Any]] = {}
    if libs_dir.exists():
        for lib_dir in libs_dir.iterdir():
            if not lib_dir.is_dir():
                continue
            per_lib_db = lib_dir / "xhs.db"
            if not per_lib_db.exists():
                continue
            try:
                uri = f"file:{per_lib_db.as_posix()}?mode=ro"
                con = sqlite3.connect(uri, uri=True)
                con.row_factory = sqlite3.Row
                try:
                    has = con.execute(
                        "SELECT name FROM sqlite_master"
                        " WHERE type='table' AND name='studio_external_reports'"
                    ).fetchone()
                    if not has:
                        continue
                    where = "WHERE (project_id = ? OR project_id IS NULL)"
                    args: list[Any] = [pid]
                    if library_id:
                        where += " AND (library_id = ? OR library_id IS NULL)"
                        args.append(library_id)
                    rows = list(con.execute(
                        f"SELECT report_id, project_id, library_id, name, source, format,"
                        f" LENGTH(content) AS content_chars, uploaded_at"
                        f" FROM studio_external_reports {where}",
                        args,
                    ))
                    for r in rows:
                        d = dict(r)
                        # report_id 是 PK，跨库不会冲突；后到的覆盖前到无所谓
                        seen[d["report_id"]] = d
                finally:
                    con.close()
            except Exception:
                continue
    items = sorted(seen.values(), key=lambda d: -(d.get("uploaded_at") or 0))
    return items


def get_external_report(report_id: str) -> dict[str, Any] | None:
    """v0.61.5 ：跨所有 per-lib .db 找该 report_id。"""
    import sqlite3
    from .. import config

    db.apply_migrations(verbose=False)
    libs_dir = config.DATA_DIR / "libraries"
    if not libs_dir.exists():
        return None
    for lib_dir in libs_dir.iterdir():
        if not lib_dir.is_dir():
            continue
        per_lib_db = lib_dir / "xhs.db"
        if not per_lib_db.exists():
            continue
        try:
            uri = f"file:{per_lib_db.as_posix()}?mode=ro"
            con = sqlite3.connect(uri, uri=True)
            con.row_factory = sqlite3.Row
            try:
                has = con.execute(
                    "SELECT name FROM sqlite_master"
                    " WHERE type='table' AND name='studio_external_reports'"
                ).fetchone()
                if not has:
                    continue
                row = con.execute(
                    "SELECT * FROM studio_external_reports WHERE report_id = ?",
                    (report_id,),
                ).fetchone()
                if row:
                    return dict(row)
            finally:
                con.close()
        except Exception:
            continue
    return None


def delete_external_report(report_id: str) -> bool:
    """v0.61.5 ：在所有 per-lib .db 里都试一次删除，命中即返回。"""
    import sqlite3
    from .. import config

    db.apply_migrations(verbose=False)
    project.ensure_bootstrap()
    pid = project.active_project_id()
    libs_dir = config.DATA_DIR / "libraries"
    deleted = False
    if libs_dir.exists():
        for lib_dir in libs_dir.iterdir():
            if not lib_dir.is_dir():
                continue
            per_lib_db = lib_dir / "xhs.db"
            if not per_lib_db.exists():
                continue
            try:
                con = sqlite3.connect(per_lib_db)
                con.row_factory = sqlite3.Row
                try:
                    has = con.execute(
                        "SELECT name FROM sqlite_master"
                        " WHERE type='table' AND name='studio_external_reports'"
                    ).fetchone()
                    if not has:
                        continue
                    cur = con.execute(
                        "DELETE FROM studio_external_reports"
                        " WHERE report_id = ? AND (project_id = ? OR project_id IS NULL)",
                        (report_id, pid),
                    )
                    if cur.rowcount > 0:
                        deleted = True
                    con.commit()
                finally:
                    con.close()
            except Exception:
                continue
    if deleted:
        try:
            from . import pipeline as _ip
            _ip.invalidate_ref_block_cache()
        except Exception:
            pass
    return deleted


# ---- Integration via gpt-4o ---------------------------------------------

INTEGRATION_SYSTEM = """\
你是「报告整合主编」。用户扔给你 N 份分析报告（可能来自不同 AI、咨询稿、自己写的笔记等），可能还有本工具自动出的双 AI 共识。

**关键 ：你的工作是「并集」不是「交集」**。每份报告都有它独到的视角和洞察，**绝对不要把不同的观点压扁成共识**。具体规则 ：

1. **每份报告的核心论点都要保留**，不管其它报告有没有同样的观点。
2. 多份报告说同一件事 → 在 evidence 里标「报告 A + 报告 B 都强调」，但不要因此降权 ；恰恰是值得加强的信号。
3. **相互矛盾的观点都要保留**，分别列出。让用户看到「报告 A 说 X、报告 B 说 Y」的分歧本身就是有价值的信号。
4. consensus_findings 应该是**全集 union 后的 findings 列表**，至少 6-10 条，不是「都同意的才进」。
5. single_side_views 是补充字段，用来标「只在一份里出现的独到观点」，至少 4-6 条。
6. **不要丢失数字 / 引用 / 案例**。每份报告里出现的关键数据点都要进 findings 的 evidence。

输出 JSON：
{
  "title": "<整合后的报告标题>",
  "executive_summary": "<6-10 句话，把所有报告的核心论点都覆盖一遍，不要压缩成 3 句>",
  "launch_mode": {
    "recommendation": "cold_start" | "hot_start" | "hybrid",
    "rationale": "<综合所有报告的理由 — 如果报告之间不一致，明说>",
    "first_week_plan": "<第一周怎么发，融合各报告的具体建议>",
    "agreement_level": "all_agree" | "leaned" | "split"
  },
  "consensus_findings": [
    {"title": "<具体论点>", "evidence": "<引用具体哪几份报告 + 它们的数据/案例>", "implication": "..."}
  ],
  "consensus_opportunities": [
    {"opportunity": "...", "why": "<引用具体报告的数据>", "suggested_angle": "..."}
  ],
  "consensus_risks": ["<分别列每份报告提到的风险>"],
  "consensus_next_steps": ["<合集，不要去重压缩>"],
  "single_side_views": [
    {"side": "<报告名>", "point": "<该报告独有的论点>", "note": "<这点为什么有价值即使其它报告没提>"}
  ],
  "charts_to_show": [
    "blue_ocean_top15" | "hook_distribution" | "timing_heatmap" |
    "top_tags" | "body_length" | "top_titles" | "comment_demand"
  ],
  "source_breakdown": [
    {"name": "<外部报告名>", "contributed": ["<它具体贡献了哪些独到观点>"]}
  ]
}

严格按 schema 输出。consensus_findings 至少 6 条（宁多勿少）。每条 evidence 都要钉到具体报告 + 具体数字/案例。
"""

_INTEGRATION_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "executive_summary": {"type": "string"},
        "launch_mode": {"type": "object"},
        "consensus_findings": {"type": "array", "items": {"type": "object"}},
        "consensus_opportunities": {"type": "array", "items": {"type": "object"}},
        "consensus_risks": {"type": "array", "items": {"type": "string"}},
        "consensus_next_steps": {"type": "array", "items": {"type": "string"}},
        "single_side_views": {"type": "array", "items": {"type": "object"}},
        "charts_to_show": {"type": "array", "items": {"type": "string"}},
        "source_breakdown": {"type": "array", "items": {"type": "object"}},
    },
    "required": ["consensus_findings", "executive_summary"],
}


async def integrate(
    source_ids: list[str],
    *,
    library_id: str | None = None,
    include_consensus_report_id: str | None = None,
    # v0.51: defaulted back to gpt-4o (Claude too expensive for daily use).
    model_spec: str = "openai:gpt-4o",
) -> dict[str, Any]:
    """Fuse external reports + (optionally) one tool-generated consensus.

    Persists an integrated_reports row and returns the full record.
    """
    db.apply_migrations(verbose=False)
    project.ensure_bootstrap()
    pid = project.active_project_id()

    # Gather external bodies.
    ext_blocks: list[str] = []
    ext_names: list[str] = []
    for sid in source_ids:
        r = get_external_report(sid)
        if not r:
            continue
        ext_names.append(r["name"])
        body = r["content"]
        ext_blocks.append(f"━━━━━ 外部报告 ：《{r['name']}》"
                          f" (来源 ：{r.get('source','—')}, 格式 ：{r.get('format','text')})\n{body}")
    if not ext_blocks:
        raise ValueError("no valid external reports referenced")

    # Optionally append the tool's own consensus for fusion.
    tool_consensus_block = ""
    if include_consensus_report_id:
        from . import pipeline as insight_pipeline
        rep = insight_pipeline.get_report(include_consensus_report_id)
        if rep and rep.get("consensus"):
            tool_consensus_block = (
                "\n\n━━━━━ 本工具自动生成的双 AI 共识报告 (供融合参考)\n"
                + json.dumps(rep["consensus"], ensure_ascii=False, indent=2)
            )

    user_msg = (
        f"共有 {len(ext_blocks)} 份外部报告需要整合"
        + (f" + 工具自身的双 AI 共识报告" if tool_consensus_block else "")
        + " ：\n\n"
        + "\n\n".join(ext_blocks)
        + tool_consensus_block
        + "\n\n请按 system 的 schema 输出整合后的统一共识报告。"
    )

    integrated_id = "int_" + uuid.uuid4().hex[:14]
    now = int(time.time())
    with db.connect() as con:
        con.execute(
            "INSERT INTO studio_integrated_reports"
            " (integrated_id, project_id, library_id, created_at, status,"
            "  source_ids, include_consensus_report_id)"
            " VALUES (?, ?, ?, ?, 'pending', ?, ?)",
            (integrated_id, pid, library_id, now,
             json.dumps(source_ids), include_consensus_report_id),
        )

    t0 = time.time()
    try:
        gen = registry.build(model_spec)[0]
        consensus = await call_for_json(
            gen, INTEGRATION_SYSTEM, user_msg,
            max_tokens=6500,
            tool_name="submit_integrated_consensus",
            schema=_INTEGRATION_SCHEMA,
        )
        elapsed = int(time.time() - t0)
        with db.connect() as con:
            con.execute(
                "UPDATE studio_integrated_reports SET status='completed',"
                " consensus_json=?, elapsed_s=? WHERE integrated_id=?",
                (json.dumps(consensus, ensure_ascii=False), elapsed, integrated_id),
            )
        try:
            from . import pipeline as _ip
            _ip.invalidate_ref_block_cache()
        except Exception:
            pass
        return {
            "integrated_id": integrated_id, "project_id": pid, "library_id": library_id,
            "created_at": now, "status": "completed", "elapsed_s": elapsed,
            "source_ids": source_ids,
            "include_consensus_report_id": include_consensus_report_id,
            "source_names": ext_names,
            "consensus": consensus,
        }
    except Exception as e:
        with db.connect() as con:
            con.execute(
                "UPDATE studio_integrated_reports SET status='failed', error=?"
                " WHERE integrated_id=?",
                (repr(e), integrated_id),
            )
        raise


def get_integrated_report(integrated_id: str) -> dict[str, Any] | None:
    db.apply_migrations(verbose=False)
    with db.connect(read_only=True) as con:
        row = con.execute(
            "SELECT * FROM studio_integrated_reports WHERE integrated_id = ?",
            (integrated_id,),
        ).fetchone()
    if not row:
        return None
    d = dict(row)
    d["source_ids"] = json.loads(d.get("source_ids") or "[]")
    d["consensus"] = json.loads(d.get("consensus_json") or "null")
    d.pop("consensus_json", None)
    return d


def list_integrated_reports(library_id: str | None = None) -> list[dict[str, Any]]:
    db.apply_migrations(verbose=False)
    project.ensure_bootstrap()
    pid = project.active_project_id()
    where = "WHERE (project_id = ? OR project_id IS NULL)"
    args: list[Any] = [pid]
    if library_id:
        where += " AND (library_id = ? OR library_id IS NULL)"
        args.append(library_id)
    with db.connect(read_only=True) as con:
        rows = list(con.execute(
            f"SELECT integrated_id, project_id, library_id, created_at, status,"
            f" source_ids, elapsed_s, error FROM studio_integrated_reports"
            f" {where} ORDER BY created_at DESC",
            args,
        ))
    out = []
    for r in rows:
        d = dict(r)
        d["source_ids"] = json.loads(d.get("source_ids") or "[]")
        out.append(d)
    return out


def latest_integrated_for_current_library() -> dict[str, Any] | None:
    """Most recent successful integrated report — used by Strategy / Composer
    to include it in prompts alongside (or instead of) the tool's own consensus.
    """
    db.apply_migrations(verbose=False)
    project.ensure_bootstrap()
    pid = project.active_project_id()
    from .. import library as _library
    lib = _library.active_lib_id()
    with db.connect(read_only=True) as con:
        # Match on library if set; also accept library_id=NULL (global integration).
        row = con.execute(
            "SELECT consensus_json FROM studio_integrated_reports"
            " WHERE (project_id = ? OR project_id IS NULL)"
            " AND (library_id = ? OR library_id IS NULL)"
            " AND status = 'completed'"
            " ORDER BY created_at DESC LIMIT 1",
            (pid, lib),
        ).fetchone()
    if not row or not row["consensus_json"]:
        return None
    try:
        return json.loads(row["consensus_json"])
    except json.JSONDecodeError:
        return None
